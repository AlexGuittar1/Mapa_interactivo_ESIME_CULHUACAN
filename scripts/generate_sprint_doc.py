from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sprint_doc():
    document = Document()

    # Title
    title = document.add_heading('Planificación del Sprint #2: Navegación Avanzada y Personalización', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project Info
    p = document.add_paragraph()
    p.add_run('Proyecto: ').bold = True
    p.add_run('Mapa Interactivo ESIME Culhuacán\n')
    p.add_run('Duración: ').bold = True
    p.add_run('2 Semanas\n')
    p.add_run('Sprint Goal: ').bold = True
    p.add_run('Implementar un sistema de navegación visual completo y herramientas de personalización para el usuario.')

    document.add_paragraph('---')

    document.add_paragraph('Este documento detalla la estructura del Sprint #2 siguiendo la metodología Scrum, enfocado en entregar un incremento de producto funcional y de alto valor para los estudiantes de la ESIME.')

    # 1. Reunión de Planificación
    document.add_heading('1. Reunión de Planificación (Sprint Planning)', level=1)
    document.add_paragraph('En esta etapa definimos el Sprint Goal y el Sprint Backlog. Basándonos en la arquitectura actual (React + Flask) y las necesidades detectadas, el enfoque será la interacción avanzada con el mapa.')

    document.add_heading('Objetivo del Sprint (Sprint Goal)', level=2)
    quote = document.add_paragraph('"Permitir al usuario no solo ver el mapa, sino interactuar con él: buscar lugares específicos, trazar rutas visuales entre puntos y personalizar su experiencia guardando sus ubicaciones favoritas."')
    quote.italic = True

    document.add_heading('Sprint Backlog (Tareas Priorizadas)', level=2)
    document.add_paragraph('Estas son las historias de usuario y tareas técnicas seleccionadas para este ciclo:')

    document.add_heading('A. Módulo de Navegación y Rutas', level=3)
    document.add_paragraph('1. Interfaz de Rutas (UI): Implementar inputs para seleccionar "Origen" y "Destino" dinámicamente.', style='List Bullet')
    document.add_paragraph('2. Visualización de Rutas (Polyline): Dibujar líneas visuales en el mapa que conecten el punto A con el punto B usando coordenadas reales de la ESIME.', style='List Bullet')
    document.add_paragraph('3. Lógica de "Tu Ubicación": Permitir cambiar el punto de origen predeterminado ("Tu Ubicación") a un punto seleccionado manualmente.', style='List Bullet')

    document.add_heading('B. Buscador y Filtros (Search & Discovery)', level=3)
    document.add_paragraph('4. Barra de Búsqueda Inteligente: Implementar autocompletado para buscar salones (ej. "Laboratorio de Computación", "Edificio 2").', style='List Bullet')
    document.add_paragraph('5. Filtros Rápidos (Chips): Botones para categorías comunes: "Cafetería", "Baños", "Gym", "Gobierno". Al hacer clic, deben centrar el mapa y mostrar info.', style='List Bullet')

    document.add_heading('C. Personalización (Guardados y Pines)', level=3)
    document.add_paragraph('6. Gestión de Favoritos: Funcionalidad para guardar edificios oficiales en una lista de "Lugares Guardados" y poder eliminarlos.', style='List Bullet')
    document.add_paragraph('7. Pines Personalizados: Permitir al usuario hacer clic largo en el mapa para colocar un pin personal, nombrarlo y guardarlo.', style='List Bullet')
    document.add_paragraph('8. Privacidad de Datos: Asegurar que los pines personales solo sean visibles para el usuario que los creó (Backend filtering por boleta).', style='List Bullet')

    document.add_heading('D. Geolocalización', level=3)
    document.add_paragraph('9. Brújula y Centrado: Botón para re-centrar el mapa en la ubicación GPS del usuario.', style='List Bullet')
    document.add_paragraph('10. Alerta de Rango: Mostrar aviso si el usuario está fuera del campus ("Estás lejos de la ESIME").', style='List Bullet')

    document.add_paragraph('---')

    # 2. Daily Scrum
    document.add_heading('2. El Scrum Diario (Daily Stand-up)', level=1)
    document.add_paragraph('Para mantener el ritmo y detectar bloqueos rápidamente, se establece la siguiente dinámica diaria:')
    
    document.add_paragraph('Formato: Reunión de 15 minutos (sincronización rápida).', style='List Bullet')
    document.add_paragraph('Preguntas Clave:', style='List Bullet')
    
    sublist = document.add_paragraph('1. ¿Qué avancé ayer? (Ej: "Conecté el frontend con el endpoint de \'Guardar Pin\'").')
    sublist.paragraph_format.left_indent = Pt(36)
    sublist2 = document.add_paragraph('2. ¿Qué haré hoy? (Ej: "Voy a implementar el dibujo de la línea de ruta en el mapa").')
    sublist2.paragraph_format.left_indent = Pt(36)
    sublist3 = document.add_paragraph('3. ¿Tengo algún bloqueo? (Ej: "La API de rutas me devuelve coordenadas vacías en ciertos edificios").')
    sublist3.paragraph_format.left_indent = Pt(36)

    document.add_paragraph('---')

    # 3. Desarrollo
    document.add_heading('3. Desarrollo del Sprint (Ejecución)', level=1)
    document.add_paragraph('Durante estas dos semanas, el equipo de desarrollo trabajará en las tareas del Backlog respetando las reglas de Scrum: No cambios en el objetivo.')

    document.add_paragraph('Tecnologías:', style='List Bullet')
    t1 = document.add_paragraph('Frontend: Uso de react-leaflet para el trazado de políneas y marcadores interactivos. Implementación de estados locales para manejar la búsqueda instantánea.')
    t1.paragraph_format.left_indent = Pt(36)
    t2 = document.add_paragraph('Backend: Optimización de endpoints en Flask (/api/routes, /api/saved_places) para asegurar respuestas rápidas (<200ms).')
    t2.paragraph_format.left_indent = Pt(36)
    
    document.add_paragraph('Calidad: Cada funcionalidad nueva debe probarse en dispositivos móviles para asegurar que los botones sean accesibles (touch-friendly).', style='List Bullet')

    document.add_paragraph('---')

    # 4. Revisión
    document.add_heading('4. Revisión del Sprint (Sprint Review)', level=1)
    document.add_paragraph('Al finalizar las dos semanas, presentaremos el incremento "hecho" al Product Owner.')
    
    document.add_paragraph('Criterios de Aceptación para la Demo:', style='List Bullet')
    document.add_paragraph('[ ] El usuario puede escribir "Cafetería" en el buscador y el mapa se mueve ahí.', style='List Bullet')
    document.add_paragraph('[ ] Al seleccionar dos puntos, aparece una línea visual conectándolos.', style='List Bullet')
    document.add_paragraph('[ ] El usuario puede guardar un punto cualquiera del mapa y verlo en su lista lateral.', style='List Bullet')
    document.add_paragraph('[ ] Si el usuario cierra y abre la app, sus lugares guardados persisten (persistencia en BD).', style='List Bullet')

    document.add_paragraph('---')

    # 5. Retrospectiva
    document.add_heading('5. Retrospectiva del Sprint (Retrospective)', level=1)
    document.add_paragraph('Análisis final para la mejora continua del proceso.')

    document.add_paragraph('Puntos a evaluar:', style='List Bullet')
    document.add_paragraph('¿Fue realista la estimación de las tareas de trazado de rutas?', style='List Bullet')
    document.add_paragraph('¿La comunicación entre el diseño (Figma/Idea) y la implementación (React) fue clara?', style='List Bullet')
    document.add_paragraph('Acción de Mejora: Documentar mejor las coordenadas de los nodos clave para evitar errores de "ruta no encontrada".', style='List Bullet')

    document.add_paragraph('---')
    
    footer = document.add_paragraph('> Este documento fue generado basándose en las mejores prácticas de Scrum descritas en la metodología ágil.')
    footer.italic = True

    document.save('SPRINT_2_PLANIFICACION.docx')

if __name__ == "__main__":
    create_sprint_doc()
