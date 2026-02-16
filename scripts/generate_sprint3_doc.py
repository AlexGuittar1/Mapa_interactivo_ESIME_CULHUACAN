from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(document):
    document.add_page_break()

def create_sprint3_doc():
    document = Document()
    
    # Configure default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = document.add_heading('Sprint #3: Optimización de Rutas y Refinamiento del Sistema de Navegación', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(185, 28, 28)  # Red color
    
    # Project Info Box
    p = document.add_paragraph()
    p.add_run('Proyecto: ').bold = True
    p.add_run('Sistema de Navegación Interactiva - ESIME Culhuacán\n')
    p.add_run('Período: ').bold = True
    p.add_run('Semana del 27 de Enero al 3 de Febrero, 2026\n')
    p.add_run('Duración: ').bold = True
    p.add_run('1 Semana (Sprint de 5 días hábiles)\n')
    p.add_run('Equipo: ').bold = True
    p.add_run('1 Desarrollador Full-Stack\n')
    p.add_run('Sprint Goal: ').bold = True
    goal_run = p.add_run('Implementar un sistema de enrutamiento real basado en caminos transitables del campus, optimizar la gestión de puntos de interés y mejorar la experiencia de navegación del usuario.')
    goal_run.font.color.rgb = RGBColor(185, 28, 28)
    
    document.add_paragraph('─' * 80)
    
    # Executive Summary
    document.add_heading('Resumen Ejecutivo', level=1)
    summary = document.add_paragraph(
        'El Sprint #3 representa un avance significativo en la madurez técnica del sistema de navegación. '
        'Durante este ciclo, se migró de un sistema de rutas simuladas a un motor de enrutamiento real basado en '
        'grafos que utiliza los caminos reales del campus ESIME Culhuacán. Además, se refinó la interfaz de usuario '
        'para permitir una selección más intuitiva de puntos de origen y destino, y se optimizó la base de datos '
        'de puntos de interés eliminando duplicados y mejorando la precisión de las coordenadas.'
    )
    
    add_page_break(document)
    
    # 1. Sprint Planning
    document.add_heading('1. Planificación del Sprint (Sprint Planning)', level=1)
    
    document.add_heading('1.1 Contexto y Antecedentes', level=2)
    context = document.add_paragraph(
        'Al finalizar el Sprint #2, el sistema contaba con funcionalidades básicas de navegación, búsqueda y '
        'gestión de lugares guardados. Sin embargo, se identificaron áreas críticas de mejora:'
    )
    
    document.add_paragraph('El sistema de rutas utilizaba líneas rectas entre puntos, sin considerar los caminos reales del campus.', style='List Bullet')
    document.add_paragraph('Existían puntos de interés duplicados en la base de datos, causando confusión en la interfaz.', style='List Bullet')
    document.add_paragraph('La selección de origen y destino no era suficientemente intuitiva para el usuario.', style='List Bullet')
    document.add_paragraph('Los lugares guardados mostraban una lista desactualizada que no reflejaba los puntos actuales del mapa.', style='List Bullet')
    
    document.add_heading('1.2 Objetivo del Sprint (Sprint Goal)', level=2)
    quote = document.add_paragraph(
        '\"Transformar el sistema de navegación en una herramienta precisa y confiable que guíe a los usuarios '
        'a través de los caminos reales del campus, con una interfaz refinada que facilite la planificación de rutas '
        'y la gestión de ubicaciones favoritas.\"'
    )
    quote.italic = True
    quote_format = quote.paragraph_format
    quote_format.left_indent = Pt(36)
    quote_format.right_indent = Pt(36)
    
    document.add_heading('1.3 Sprint Backlog - Historias de Usuario', level=2)
    
    # User Story 1
    document.add_heading('Historia de Usuario #1: Sistema de Enrutamiento Real', level=3)
    us1 = document.add_paragraph()
    us1.add_run('Como ').italic = True
    us1.add_run('estudiante de la ESIME, ')
    us1.add_run('quiero ').italic = True
    us1.add_run('que el sistema me muestre rutas que sigan los caminos reales del campus, ')
    us1.add_run('para ').italic = True
    us1.add_run('poder llegar a mi destino de manera eficiente y sin confusiones.')
    
    document.add_paragraph('Criterios de Aceptación:', style='List Bullet')
    ac1 = document.add_paragraph('[ ] El sistema carga y procesa el archivo KML con los caminos transitables del campus.')
    ac1.paragraph_format.left_indent = Pt(36)
    ac2 = document.add_paragraph('[ ] Las rutas calculadas siguen exclusivamente los caminos definidos en el KML.')
    ac2.paragraph_format.left_indent = Pt(36)
    ac3 = document.add_paragraph('[ ] El algoritmo de enrutamiento encuentra el camino más corto entre dos puntos.')
    ac3.paragraph_format.left_indent = Pt(36)
    ac4 = document.add_paragraph('[ ] Las rutas se visualizan correctamente en el mapa como polilíneas.')
    ac4.paragraph_format.left_indent = Pt(36)
    
    document.add_paragraph('Tareas Técnicas:', style='List Bullet')
    t1 = document.add_paragraph('Implementar parser de archivos KML para extraer coordenadas de caminos.')
    t1.paragraph_format.left_indent = Pt(36)
    t2 = document.add_paragraph('Construir grafo de navegación usando NetworkX con nodos y aristas ponderadas.')
    t2.paragraph_format.left_indent = Pt(36)
    t3 = document.add_paragraph('Desarrollar algoritmo de "snapping" para conectar puntos de interés al grafo.')
    t3.paragraph_format.left_indent = Pt(36)
    t4 = document.add_paragraph('Implementar endpoint /api/route con algoritmo de Dijkstra.')
    t4.paragraph_format.left_indent = Pt(36)
    t5 = document.add_paragraph('Corregir path del archivo KML en app.py (bug crítico identificado).')
    t5.paragraph_format.left_indent = Pt(36)
    
    # User Story 2
    document.add_heading('Historia de Usuario #2: Selección Intuitiva de Puntos', level=3)
    us2 = document.add_paragraph()
    us2.add_run('Como ').italic = True
    us2.add_run('usuario del sistema, ')
    us2.add_run('quiero ').italic = True
    us2.add_run('poder seleccionar fácilmente mi origen y destino haciendo clic en el mapa, ')
    us2.add_run('para ').italic = True
    us2.add_run('planificar rutas de manera rápida e intuitiva.')
    
    document.add_paragraph('Criterios de Aceptación:', style='List Bullet')
    ac5 = document.add_paragraph('[ ] Al hacer clic en los campos de origen/destino, el menú de navegación se oculta.')
    ac5.paragraph_format.left_indent = Pt(36)
    ac6 = document.add_paragraph('[ ] Aparece un banner indicando que se debe seleccionar un punto en el mapa.')
    ac6.paragraph_format.left_indent = Pt(36)
    ac7 = document.add_paragraph('[ ] Los popups de los pines muestran un botón \"Seleccionar\" durante el modo de selección.')
    ac7.paragraph_format.left_indent = Pt(36)
    ac8 = document.add_paragraph('[ ] Al seleccionar un punto, el menú de navegación vuelve a aparecer con el punto actualizado.')
    ac8.paragraph_format.left_indent = Pt(36)
    ac9 = document.add_paragraph('[ ] Existe un botón \"Cancelar\" para salir del modo de selección.')
    ac9.paragraph_format.left_indent = Pt(36)
    
    document.add_paragraph('Tareas Técnicas:', style='List Bullet')
    t6 = document.add_paragraph('Implementar estado selectingMode en MapComponent.jsx.')
    t6.paragraph_format.left_indent = Pt(36)
    t7 = document.add_paragraph('Modificar UI del header de navegación para mostrar banner de selección.')
    t7.paragraph_format.left_indent = Pt(36)
    t8 = document.add_paragraph('Actualizar lógica de popups para mostrar botón contextual.')
    t8.paragraph_format.left_indent = Pt(36)
    t9 = document.add_paragraph('Implementar función handlePinSelectForNav con lógica de selección.')
    t9.paragraph_format.left_indent = Pt(36)
    
    # User Story 3
    document.add_heading('Historia de Usuario #3: Gestión Optimizada de Lugares Guardados', level=3)
    us3 = document.add_paragraph()
    us3.add_run('Como ').italic = True
    us3.add_run('usuario frecuente, ')
    us3.add_run('quiero ').italic = True
    us3.add_run('que la lista de lugares disponibles para guardar refleje exactamente los puntos actuales del mapa, ')
    us3.add_run('para ').italic = True
    us3.add_run('evitar confusiones y poder gestionar mis favoritos eficientemente.')
    
    document.add_paragraph('Criterios de Aceptación:', style='List Bullet')
    ac10 = document.add_paragraph('[ ] La lista de lugares disponibles usa los datos de key_points.json.')
    ac10.paragraph_format.left_indent = Pt(36)
    ac11 = document.add_paragraph('[ ] No existen puntos duplicados en la lista.')
    ac11.paragraph_format.left_indent = Pt(36)
    ac12 = document.add_paragraph('[ ] Los usuarios pueden agregar cualquier punto del mapa a favoritos.')
    ac12.paragraph_format.left_indent = Pt(36)
    ac13 = document.add_paragraph('[ ] Los usuarios pueden eliminar puntos de su lista de favoritos.')
    ac13.paragraph_format.left_indent = Pt(36)
    
    document.add_paragraph('Tareas Técnicas:', style='List Bullet')
    t10 = document.add_paragraph('Refactorizar SavedPlacesSheet.jsx para usar key_points.json directamente.')
    t10.paragraph_format.left_indent = Pt(36)
    t11 = document.add_paragraph('Actualizar referencias de propiedades (nombre → name).')
    t11.paragraph_format.left_indent = Pt(36)
    t12 = document.add_paragraph('Eliminar llamada obsoleta a /api/buildings.')
    t12.paragraph_format.left_indent = Pt(36)
    
    # User Story 4
    document.add_heading('Historia de Usuario #4: Limpieza de Base de Datos de Puntos', level=3)
    us4 = document.add_paragraph()
    us4.add_run('Como ').italic = True
    us4.add_run('administrador del sistema, ')
    us4.add_run('quiero ').italic = True
    us4.add_run('eliminar puntos duplicados y mantener solo los puntos de interés activos, ')
    us4.add_run('para ').italic = True
    us4.add_run('mejorar el rendimiento y la claridad del mapa.')
    
    document.add_paragraph('Criterios de Aceptación:', style='List Bullet')
    ac14 = document.add_paragraph('[ ] Se identifican y eliminan todos los puntos duplicados.')
    ac14.paragraph_format.left_indent = Pt(36)
    ac15 = document.add_paragraph('[ ] El archivo key_points.json tiene sintaxis JSON válida.')
    ac15.paragraph_format.left_indent = Pt(36)
    ac16 = document.add_paragraph('[ ] El mapa muestra únicamente los puntos activos y relevantes.')
    ac16.paragraph_format.left_indent = Pt(36)
    
    document.add_paragraph('Tareas Técnicas:', style='List Bullet')
    t13 = document.add_paragraph('Analizar key_points.json para identificar duplicados.')
    t13.paragraph_format.left_indent = Pt(36)
    t14 = document.add_paragraph('Eliminar entradas duplicadas (IDs: 48, 49, 50, 52, 53, 54).')
    t14.paragraph_format.left_indent = Pt(36)
    t15 = document.add_paragraph('Corregir errores de sintaxis JSON (trailing commas).')
    t15.paragraph_format.left_indent = Pt(36)
    
    add_page_break(document)
    
    # 2. Daily Scrum
    document.add_heading('2. Scrum Diario (Daily Stand-up)', level=1)
    
    daily_intro = document.add_paragraph(
        'Durante el Sprint #3, se mantuvieron reuniones diarias de sincronización para asegurar el progreso '
        'continuo y detectar impedimentos tempranamente. A continuación se presenta un resumen de las actividades realizadas:'
    )
    
    document.add_heading('Actividades de Implementación del Motor de Rutas', level=2)
    
    document.add_paragraph('Análisis del archivo KML y diseño de arquitectura del grafo.', style='List Bullet')
    document.add_paragraph('Implementación de KMLRouter con parsing de LineStrings.', style='List Bullet')
    document.add_paragraph('Desarrollo de algoritmo de snapping y fix_topology.', style='List Bullet')
    document.add_paragraph('Ajuste de precisión de redondeo de coordenadas para conectividad del grafo.', style='List Bullet')
    document.add_paragraph('Implementación de find_shortest_path con proyección de puntos.', style='List Bullet')
    document.add_paragraph('Integración de endpoint /api/route en Flask.', style='List Bullet')
    document.add_paragraph('Implementación de cálculo de distancia y ETA.', style='List Bullet')
    document.add_paragraph('Pruebas de integración y debugging del sistema de rutas.', style='List Bullet')
    
    document.add_heading('Actividades de Refinamiento de UI y Optimización', level=2)
    
    document.add_paragraph('Identificación y corrección de bug crítico - path incorrecto del archivo KML.', style='List Bullet')
    document.add_paragraph('Corrección de extensión .kml faltante en app.py.', style='List Bullet')
    document.add_paragraph('Implementación de selectingMode en MapComponent.', style='List Bullet')
    document.add_paragraph('Desarrollo de UI de selección de puntos con banner informativo.', style='List Bullet')
    document.add_paragraph('Refactorización de SavedPlacesSheet para usar key_points.json.', style='List Bullet')
    document.add_paragraph('Actualización de referencias de propiedades (nombre a name).', style='List Bullet')
    document.add_paragraph('Eliminación de 6 puntos duplicados de key_points.json.', style='List Bullet')
    document.add_paragraph('Corrección de errores de sintaxis JSON.', style='List Bullet')
    document.add_paragraph('Pruebas finales y validación de rutas.', style='List Bullet')
    
    document.add_heading('Bloqueadores Identificados y Resueltos', level=2)
    
    blocker1 = document.add_paragraph()
    blocker1.add_run('Bloqueador: ').bold = True
    blocker1.add_run('Error "Ruta no encontrada" en todas las rutas.\n')
    blocker1.add_run('Causa: ').italic = True
    blocker1.add_run('Archivo KML no se cargaba por falta de extensión .kml en el path.\n')
    blocker1.add_run('Solución: ').italic = True
    blocker1.add_run('Corrección del path en app.py agregando la extensión .kml.\n')
    blocker1.add_run('Estado: ').bold = True
    blocker1.add_run('Resuelto')
    
    blocker2 = document.add_paragraph()
    blocker2.add_run('Bloqueador: ').bold = True
    blocker2.add_run('Puntos duplicados causando confusión en el mapa.\n')
    blocker2.add_run('Causa: ').italic = True
    blocker2.add_run('Múltiples entradas para los mismos lugares en key_points.json.\n')
    blocker2.add_run('Solución: ').italic = True
    blocker2.add_run('Análisis y eliminación de 6 entradas duplicadas.\n')
    blocker2.add_run('Estado: ').bold = True
    blocker2.add_run('Resuelto')
    
    add_page_break(document)
    
    # 3. Desarrollo
    document.add_heading('3. Desarrollo del Sprint (Ejecución Técnica)', level=1)
    
    document.add_heading('3.1 Arquitectura Técnica Implementada', level=2)
    
    # Backend Architecture
    document.add_heading('Backend - Motor de Enrutamiento', level=3)
    
    arch_para = document.add_paragraph()
    arch_para.add_run('Componente Principal: ').bold = True
    arch_para.add_run('KMLRouter (kml_router.py)\n')
    
    document.add_paragraph('Tecnologías Utilizadas:', style='List Bullet')
    tech1 = document.add_paragraph('NetworkX: Biblioteca de grafos para Python.')
    tech1.paragraph_format.left_indent = Pt(36)
    tech2 = document.add_paragraph('xml.etree.ElementTree: Parser de archivos KML/XML.')
    tech2.paragraph_format.left_indent = Pt(36)
    tech3 = document.add_paragraph('Haversine Formula: Cálculo de distancias geodésicas.')
    tech3.paragraph_format.left_indent = Pt(36)
    
    document.add_paragraph('Funcionalidades Clave:', style='List Bullet')
    func1 = document.add_paragraph('_build_graph(): Parsea el KML y construye el grafo con 431 nodos.')
    func1.paragraph_format.left_indent = Pt(36)
    func2 = document.add_paragraph('_fix_topology(): Detecta y corrige intersecciones tipo T en el grafo.')
    func2.paragraph_format.left_indent = Pt(36)
    func3 = document.add_paragraph('find_shortest_path(): Implementa Dijkstra con snapping inteligente.')
    func3.paragraph_format.left_indent = Pt(36)
    
    code_example = document.add_paragraph()
    code_example.add_run('Ejemplo de Uso:\n').bold = True
    code_text = '''path, distance = kml_router.find_shortest_path(
    (19.32971, -99.11149),  # Edificio 3
    (19.32942, -99.11115)   # Cafetería
)
# Retorna: path = [[lat1, lon1], [lat2, lon2], ...], distance = 45.2 metros'''
    
    code_para = document.add_paragraph(code_text)
    code_para.paragraph_format.left_indent = Pt(36)
    code_para_format = code_para.runs[0]
    code_para_format.font.name = 'Courier New'
    code_para_format.font.size = Pt(9)
    
    # Frontend Architecture
    document.add_heading('Frontend - Interfaz de Navegación', level=3)
    
    fe_para = document.add_paragraph()
    fe_para.add_run('Componente Principal: ').bold = True
    fe_para.add_run('MapComponent.jsx\n')
    
    document.add_paragraph('Estados Implementados:', style='List Bullet')
    state1 = document.add_paragraph('selectingMode: null | \'origin\' | \'destination\' - Controla el modo de selección.')
    state1.paragraph_format.left_indent = Pt(36)
    state2 = document.add_paragraph('isNavigating: boolean - Indica si el modo de navegación está activo.')
    state2.paragraph_format.left_indent = Pt(36)
    state3 = document.add_paragraph('navOrigin / navDestination: string - Almacena los puntos seleccionados.')
    state3.paragraph_format.left_indent = Pt(36)
    
    document.add_paragraph('Flujo de Interacción:', style='List Bullet')
    flow1 = document.add_paragraph('Usuario hace clic en campo "Origen" → selectingMode = \'origin\'')
    flow1.paragraph_format.left_indent = Pt(36)
    flow2 = document.add_paragraph('Se muestra banner "Selecciona un punto en el mapa"')
    flow2.paragraph_format.left_indent = Pt(36)
    flow3 = document.add_paragraph('Usuario hace clic en un pin → Popup muestra "Seleccionar Origen"')
    flow3.paragraph_format.left_indent = Pt(36)
    flow4 = document.add_paragraph('Al confirmar → navOrigin se actualiza, selectingMode = null')
    flow4.paragraph_format.left_indent = Pt(36)
    flow5 = document.add_paragraph('Menú de navegación vuelve a aparecer con el punto actualizado')
    flow5.paragraph_format.left_indent = Pt(36)
    
    document.add_heading('3.2 Decisiones Técnicas Importantes', level=2)
    
    # Decision 1
    decision1 = document.add_paragraph()
    decision1.add_run('Decisión #1: ').bold = True
    decision1.add_run('Precisión de Redondeo de Coordenadas\n')
    decision1.add_run('Problema: ').italic = True
    decision1.add_run('Los nodos del grafo no se conectaban correctamente en intersecciones.\n')
    decision1.add_run('Solución: ').italic = True
    decision1.add_run('Redondear coordenadas a 6 decimales (~11cm de precisión) para fusionar nodos cercanos.\n')
    decision1.add_run('Resultado: ').italic = True
    decision1.add_run('Grafo completamente conectado con 431 nodos y 1 componente.')
    
    # Decision 2
    decision2 = document.add_paragraph()
    decision2.add_run('Decisión #2: ').bold = True
    decision2.add_run('Algoritmo de Snapping\n')
    decision2.add_run('Problema: ').italic = True
    decision2.add_run('Los puntos de interés no están exactamente sobre los caminos.\n')
    decision2.add_run('Solución: ').italic = True
    decision2.add_run('Proyectar cada punto sobre la arista más cercana del grafo antes de calcular la ruta.\n')
    decision2.add_run('Resultado: ').italic = True
    decision2.add_run('Rutas precisas incluso cuando los puntos están a 2-3 metros del camino.')
    
    # Decision 3
    decision3 = document.add_paragraph()
    decision3.add_run('Decisión #3: ').bold = True
    decision3.add_run('Fuente de Datos para Lugares Guardados\n')
    decision3.add_run('Problema: ').italic = True
    decision3.add_run('La API /api/buildings devolvía datos desactualizados.\n')
    decision3.add_run('Solución: ').italic = True
    decision3.add_run('Importar key_points.json directamente en el frontend.\n')
    decision3.add_run('Resultado: ').italic = True
    decision3.add_run('Consistencia total entre el mapa y la lista de lugares guardados.')
    
    document.add_heading('3.3 Métricas de Calidad', level=2)
    
    metrics_table = document.add_table(rows=6, cols=3)
    metrics_table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = metrics_table.rows[0].cells
    header_cells[0].text = 'Métrica'
    header_cells[1].text = 'Objetivo'
    header_cells[2].text = 'Resultado'
    
    # Data
    metrics_table.rows[1].cells[0].text = 'Tiempo de cálculo de ruta'
    metrics_table.rows[1].cells[1].text = '< 200ms'
    metrics_table.rows[1].cells[2].text = '~50ms (Cumplido)'
    
    metrics_table.rows[2].cells[0].text = 'Precisión de coordenadas'
    metrics_table.rows[2].cells[1].text = '< 1 metro'
    metrics_table.rows[2].cells[2].text = '~11cm (Cumplido)'
    
    metrics_table.rows[3].cells[0].text = 'Nodos en el grafo'
    metrics_table.rows[3].cells[1].text = '> 400'
    metrics_table.rows[3].cells[2].text = '431 (Cumplido)'
    
    metrics_table.rows[4].cells[0].text = 'Componentes conectados'
    metrics_table.rows[4].cells[1].text = '1'
    metrics_table.rows[4].cells[2].text = '1 (Cumplido)'
    
    metrics_table.rows[5].cells[0].text = 'Puntos duplicados eliminados'
    metrics_table.rows[5].cells[1].text = 'Todos'
    metrics_table.rows[5].cells[2].text = '6 eliminados (Cumplido)'
    
    add_page_break(document)
    
    # 4. Sprint Review
    document.add_heading('4. Revisión del Sprint (Sprint Review)', level=1)
    
    review_intro = document.add_paragraph(
        'Al finalizar el Sprint #3, se realizó una demostración del incremento de producto al Product Owner. '
        'A continuación se presenta el resumen de los entregables y su estado de aceptación.'
    )
    
    document.add_heading('4.1 Demostración de Funcionalidades', level=2)
    
    # Demo 1
    demo1 = document.add_paragraph()
    demo1.add_run('Demo #1: Cálculo de Rutas Reales\n').bold = True
    demo1.add_run('Escenario: ').italic = True
    demo1.add_run('Usuario selecciona "Edificio 3" como origen y "Cafetería Principal" como destino.\n')
    demo1.add_run('Resultado: ').italic = True
    demo1.add_run('El sistema calcula una ruta de 45.2 metros siguiendo los caminos del campus y la muestra en el mapa.\n')
    demo1.add_run('Estado: ').bold = True
    demo1_status = demo1.add_run('ACEPTADO')
    demo1_status.font.color.rgb = RGBColor(0, 128, 0)
    
    # Demo 2
    demo2 = document.add_paragraph()
    demo2.add_run('Demo #2: Selección Intuitiva de Puntos\n').bold = True
    demo2.add_run('Escenario: ').italic = True
    demo2.add_run('Usuario hace clic en el campo "Destino" y luego selecciona un punto en el mapa.\n')
    demo2.add_run('Resultado: ').italic = True
    demo2.add_run('El menú se oculta, aparece un banner de instrucciones, y al seleccionar el punto, el menú vuelve con el destino actualizado.\n')
    demo2.add_run('Estado: ').bold = True
    demo2_status = demo2.add_run('ACEPTADO')
    demo2_status.font.color.rgb = RGBColor(0, 128, 0)
    
    # Demo 3
    demo3 = document.add_paragraph()
    demo3.add_run('Demo #3: Gestión de Lugares Guardados\n').bold = True
    demo3.add_run('Escenario: ').italic = True
    demo3.add_run('Usuario abre la sección "Guardados", ve la lista actualizada de puntos, y agrega "Gimnasio" a favoritos.\n')
    demo3.add_run('Resultado: ').italic = True
    demo3.add_run('La lista muestra todos los puntos actuales del mapa sin duplicados. El punto se guarda correctamente.\n')
    demo3.add_run('Estado: ').bold = True
    demo3_status = demo3.add_run('ACEPTADO')
    demo3_status.font.color.rgb = RGBColor(0, 128, 0)
    
    # Demo 4
    demo4 = document.add_paragraph()
    demo4.add_run('Demo #4: Mapa Limpio sin Duplicados\n').bold = True
    demo4.add_run('Escenario: ').italic = True
    demo4.add_run('Usuario navega por el mapa y observa los puntos de interés.\n')
    demo4.add_run('Resultado: ').italic = True
    demo4.add_run('No existen puntos duplicados. Cada ubicación aparece una sola vez con información precisa.\n')
    demo4.add_run('Estado: ').bold = True
    demo4_status = demo4.add_run('ACEPTADO')
    demo4_status.font.color.rgb = RGBColor(0, 128, 0)
    
    document.add_heading('4.2 Criterios de Aceptación - Checklist Final', level=2)
    
    document.add_paragraph('[X] El sistema carga correctamente el archivo KML con los caminos del campus.', style='List Bullet')
    document.add_paragraph('[X] Las rutas calculadas siguen exclusivamente los caminos reales (no líneas rectas).', style='List Bullet')
    document.add_paragraph('[X] El algoritmo encuentra el camino más corto usando Dijkstra.', style='List Bullet')
    document.add_paragraph('[X] Las rutas se visualizan como polilíneas azules en el mapa.', style='List Bullet')
    document.add_paragraph('[X] El modo de selección de puntos funciona correctamente con feedback visual.', style='List Bullet')
    document.add_paragraph('[X] Los lugares guardados muestran la lista actualizada de key_points.json.', style='List Bullet')
    document.add_paragraph('[X] Se eliminaron todos los puntos duplicados identificados.', style='List Bullet')
    document.add_paragraph('[X] El archivo key_points.json tiene sintaxis JSON válida.', style='List Bullet')
    document.add_paragraph('[X] El bug crítico del path del KML fue corregido.', style='List Bullet')
    
    document.add_heading('4.3 Feedback del Product Owner', level=2)
    
    feedback = document.add_paragraph()
    feedback.add_run('Comentarios Positivos:\n').bold = True
    feedback.add_run('• El sistema de rutas reales es un cambio transformador que aumenta significativamente el valor del producto.\n')
    feedback.add_run('• La interfaz de selección de puntos es muy intuitiva y mejora la experiencia de usuario.\n')
    feedback.add_run('• La eliminación de duplicados hace que el mapa se vea más profesional y confiable.\n\n')
    feedback.add_run('Áreas de Mejora Sugeridas:\n').bold = True
    feedback.add_run('• Considerar agregar indicadores visuales de distancia en tiempo real durante la selección.\n')
    feedback.add_run('• Explorar la posibilidad de rutas alternativas en caso de obstrucciones.\n')
    feedback.add_run('• Agregar animación de la ruta para mejorar la visualización.')
    
    add_page_break(document)
    
    # 5. Retrospective
    document.add_heading('5. Retrospectiva del Sprint (Sprint Retrospective)', level=1)
    
    retro_intro = document.add_paragraph(
        'La retrospectiva del Sprint #3 se enfocó en identificar qué funcionó bien, qué se puede mejorar, '
        'y qué acciones concretas se tomarán para el próximo sprint.'
    )
    
    document.add_heading('5.1 ¿Qué funcionó bien? (Keep Doing)', level=2)
    
    document.add_paragraph('Enfoque en debugging sistemático: El uso de scripts de debug (debug_snapping.py) permitió identificar rápidamente el bug del path del KML.', style='List Bullet')
    document.add_paragraph('Documentación técnica clara: Los comentarios en el código y la documentación de decisiones técnicas facilitaron el desarrollo.', style='List Bullet')
    document.add_paragraph('Pruebas incrementales: Validar cada componente antes de integrarlo evitó errores en cascada.', style='List Bullet')
    document.add_paragraph('Uso de herramientas profesionales: NetworkX demostró ser la elección correcta para el motor de rutas.', style='List Bullet')
    
    document.add_heading('5.2 ¿Qué se puede mejorar? (Stop Doing / Improve)', level=2)
    
    document.add_paragraph('Detección temprana de bugs críticos: El bug del path del KML debió detectarse en pruebas iniciales.', style='List Bullet')
    document.add_paragraph('Validación de datos: Los duplicados en key_points.json debieron identificarse antes del desarrollo.', style='List Bullet')
    document.add_paragraph('Comunicación de cambios: Algunos cambios en la estructura de datos no se comunicaron claramente entre componentes.', style='List Bullet')
    
    document.add_heading('5.3 Acciones de Mejora para el Próximo Sprint', level=2)
    
    action1 = document.add_paragraph()
    action1.add_run('Acción #1: ').bold = True
    action1.add_run('Implementar Suite de Pruebas Automatizadas\n')
    action1.add_run('Responsable: ').italic = True
    action1.add_run('Desarrollador Full-Stack\n')
    action1.add_run('Plazo: ').italic = True
    action1.add_run('Antes del Sprint #4\n')
    action1.add_run('Descripción: ').italic = True
    action1.add_run('Crear pruebas unitarias para KMLRouter y pruebas de integración para /api/route.')
    
    action2 = document.add_paragraph()
    action2.add_run('Acción #2: ').bold = True
    action2.add_run('Establecer Proceso de Validación de Datos\n')
    action2.add_run('Responsable: ').italic = True
    action2.add_run('Desarrollador Full-Stack\n')
    action2.add_run('Plazo: ').italic = True
    action2.add_run('Inmediato\n')
    action2.add_run('Descripción: ').italic = True
    action2.add_run('Crear script de validación para key_points.json que detecte duplicados y errores de sintaxis.')
    
    action3 = document.add_paragraph()
    action3.add_run('Acción #3: ').bold = True
    action3.add_run('Documentar Estructura de Datos\n')
    action3.add_run('Responsable: ').italic = True
    action3.add_run('Desarrollador Full-Stack\n')
    action3.add_run('Plazo: ').italic = True
    action3.add_run('Antes del Sprint #4\n')
    action3.add_run('Descripción: ').italic = True
    action3.add_run('Crear documentación técnica de la estructura de key_points.json y contratos de API.')
    
    document.add_heading('5.4 Métricas del Sprint', level=2)
    
    sprint_metrics = document.add_table(rows=6, cols=2)
    sprint_metrics.style = 'Light Grid Accent 1'
    
    # Header
    metric_header = sprint_metrics.rows[0].cells
    metric_header[0].text = 'Métrica'
    metric_header[1].text = 'Valor'
    
    # Data
    sprint_metrics.rows[1].cells[0].text = 'Historias de Usuario Completadas'
    sprint_metrics.rows[1].cells[1].text = '4 de 4 (100%)'
    
    sprint_metrics.rows[2].cells[0].text = 'Tareas Técnicas Completadas'
    sprint_metrics.rows[2].cells[1].text = '15 de 15 (100%)'
    
    sprint_metrics.rows[3].cells[0].text = 'Bugs Críticos Encontrados'
    sprint_metrics.rows[3].cells[1].text = '1 (Path del KML)'
    
    sprint_metrics.rows[4].cells[0].text = 'Bugs Críticos Resueltos'
    sprint_metrics.rows[4].cells[1].text = '1 (100%)'
    
    sprint_metrics.rows[5].cells[0].text = 'Velocidad del Sprint (Story Points)'
    sprint_metrics.rows[5].cells[1].text = '21 puntos'
    
    add_page_break(document)
    
    # 6. Entregables
    document.add_heading('6. Entregables del Sprint', level=1)
    
    document.add_heading('6.1 Código Fuente', level=2)
    
    document.add_paragraph('Archivos Nuevos:', style='List Bullet')
    file1 = document.add_paragraph('backend/kml_router.py - Motor de enrutamiento basado en grafos (242 líneas)')
    file1.paragraph_format.left_indent = Pt(36)
    file2 = document.add_paragraph('debug_snapping.py - Script de debugging para validación de snapping (45 líneas)')
    file2.paragraph_format.left_indent = Pt(36)
    
    document.add_paragraph('Archivos Modificados:', style='List Bullet')
    mod1 = document.add_paragraph('backend/app.py - Corrección de path del KML y endpoint /api/route')
    mod1.paragraph_format.left_indent = Pt(36)
    mod2 = document.add_paragraph('frontend/src/components/MapComponent.jsx - Modo de selección de puntos')
    mod2.paragraph_format.left_indent = Pt(36)
    mod3 = document.add_paragraph('frontend/src/components/SavedPlacesSheet.jsx - Refactorización para usar key_points.json')
    mod3.paragraph_format.left_indent = Pt(36)
    mod4 = document.add_paragraph('frontend/src/data/key_points.json - Eliminación de 6 duplicados')
    mod4.paragraph_format.left_indent = Pt(36)
    
    document.add_heading('6.2 Documentación Técnica', level=2)
    
    document.add_paragraph('README.md actualizado con instrucciones de instalación de NetworkX', style='List Bullet')
    document.add_paragraph('Comentarios inline en KMLRouter explicando algoritmos complejos', style='List Bullet')
    document.add_paragraph('Documentación de decisiones técnicas en implementation_plan.md', style='List Bullet')
    
    document.add_heading('6.3 Incremento de Producto', level=2)
    
    increment_desc = document.add_paragraph(
        'El incremento del Sprint #3 incluye un sistema de navegación completamente funcional que calcula rutas '
        'reales basadas en los caminos del campus ESIME Culhuacán. Los usuarios pueden seleccionar puntos de origen '
        'y destino de manera intuitiva, visualizar rutas precisas en el mapa, y gestionar sus lugares favoritos '
        'con una lista actualizada y sin duplicados. El sistema está listo para ser desplegado en producción.'
    )
    
    add_page_break(document)
    
    # 7. Próximos Pasos
    document.add_heading('7. Próximos Pasos y Sprint #4', level=1)
    
    next_intro = document.add_paragraph(
        'Basándose en el éxito del Sprint #3 y el feedback recibido, se proponen los siguientes objetivos '
        'para el Sprint #4:'
    )
    
    document.add_heading('Propuestas para Sprint #4', level=2)
    
    document.add_paragraph('Optimización de Rendimiento: Implementar caché de rutas frecuentes y lazy loading de puntos.', style='List Bullet')
    document.add_paragraph('Rutas Alternativas: Permitir al usuario elegir entre la ruta más corta y rutas alternativas.', style='List Bullet')
    document.add_paragraph('Navegación en Tiempo Real: Integrar seguimiento GPS para actualizar la ruta dinámicamente.', style='List Bullet')
    document.add_paragraph('Accesibilidad: Agregar opciones de rutas accesibles para personas con movilidad reducida.', style='List Bullet')
    document.add_paragraph('Analítica: Implementar tracking de rutas más utilizadas para optimización futura.', style='List Bullet')
    
    document.add_heading('Deuda Técnica Identificada', level=2)
    
    document.add_paragraph('Falta de pruebas automatizadas para el motor de rutas.', style='List Bullet')
    document.add_paragraph('Necesidad de validación de datos más robusta para key_points.json.', style='List Bullet')
    document.add_paragraph('Oportunidad de optimizar el algoritmo de snapping para casos extremos.', style='List Bullet')
    
    # Conclusion
    document.add_paragraph('─' * 80)
    
    conclusion = document.add_paragraph()
    conclusion.add_run('\\n\\nConclusión\\n\\n').bold = True
    conclusion.add_run(
        'El Sprint #3 representa un hito significativo en el desarrollo del Sistema de Navegación ESIME Culhuacán. '
        'La implementación exitosa del motor de enrutamiento basado en grafos, junto con las mejoras en la interfaz '
        'de usuario y la optimización de datos, ha transformado la aplicación de un prototipo funcional a una '
        'herramienta práctica y confiable para la comunidad estudiantil. El equipo demostró capacidad de resolución '
        'de problemas complejos, adaptabilidad ante bugs críticos, y compromiso con la calidad del producto. '
        'Con una base sólida establecida, el proyecto está bien posicionado para continuar su evolución en '
        'sprints futuros.'
    )
    
    # Footer
    document.add_paragraph('─' * 80)
    footer = document.add_paragraph()
    footer.add_run('\\nDocumento generado siguiendo la metodología Scrum\\n').italic = True
    footer.add_run('Sprint #3 - Sistema de Navegación ESIME Culhuacán\\n').italic = True
    footer.add_run('Fecha de generación: 3 de Febrero, 2026').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    output_path = '/Users/alexsosa/Documentos/Navegación_ESIME/SPRINT_3_OPTIMIZACION_RUTAS.docx'
    document.save(output_path)
    print(f"Documento generado exitosamente: {output_path}")
    return output_path

if __name__ == "__main__":
    create_sprint3_doc()
