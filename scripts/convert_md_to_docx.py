import docx
import re
import os

def parse_markdown_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Horizontal Rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 40)
            
        # Lists
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            format_text(p, text)
            
        # Numbered Lists (Simple detection)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', line)
            format_text(p, text)
            
        # Normal Text
        else:
            p = doc.add_paragraph()
            format_text(p, line)
            
    doc.save(docx_path)
    print(f"Document saved to {docx_path}")

def format_text(paragraph, text):
    # Handle bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
             # Basic handling for inline code if needed, otherwise just text
             run = paragraph.add_run(part)
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    md_file = 'SPRINT_REVIEW.md'
    docx_file = 'SPRINT_REVIEW.docx'
    
    if os.path.exists(md_file):
        parse_markdown_to_docx(md_file, docx_file)
    else:
        print(f"File {md_file} not found.")
